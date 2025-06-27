import os
from pathlib import Path
from typing import List
import pypdf
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def process_file(self, file_path: str, file_name: str) -> List[LangchainDocument]:
        """Process a single file and return chunks"""
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            text = self._extract_from_pdf(file_path)
        elif file_extension == '.txt':
            text = self._extract_from_txt(file_path)
        elif file_extension == '.docx':
            text = self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create LangChain documents
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    'source': file_name,
                    'chunk_index': i,
                    'file_type': file_extension
                }
            )
            documents.append(doc)
        
        return documents
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text